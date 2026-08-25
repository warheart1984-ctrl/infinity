"""Build the 15-chapter constitutional enforcement stack docx."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def p(t, bold=False): 
    para = doc.add_paragraph()
    run = para.add_run(t); run.bold = bold
    return para
def bullets(items):
    for it in items: doc.add_paragraph(it, style="List Bullet")
def code(t):
    para = doc.add_paragraph()
    run = para.add_run(t); run.font.name = "Consolas"; run.font.size = Pt(9)
    para.paragraph_format.left_indent = Inches(0.3)

# Title page
title = doc.add_heading("The Constitutional Enforcement Stack", level=0)
sub = doc.add_paragraph("AI Organism Governance — From Receipts to Attested Execution")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta = doc.add_paragraph("Project Infinity / AAES-OS — August 2026\nCheckpoints: b9852d7 (port), caab105 (stack + laws), 79bdbb6 (dominance), e06c823 (14 sinks), bebdf1c (trust-root + attestation)")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

h1("Chapter 1 — Executive Summary")
p("This document consolidates the constitutional enforcement work built across the AI Organism (Project Infinity / AAIS) and the AAES-OS TypeScript spine. What began as a receipt-parity port became a full authority boundary: the organism can now seal evidence, judge its own mutations against canonical invariants, refuse unattested execution instances, and dominate every law-state write path in the codebase.")
bullets([
    "Cross-runtime determinism: Node (Sovereign-X) and Python issue byte-identical receipt ids.",
    "Sink-level dominance: all 14 law registries refuse writes without CEN approval bound to the exact record.",
    "INV-021 hard gate: law mutations require a VT authority token — no valid token, no state transition.",
    "TOCTOU closure: commits execute only CEN-frozen approved objects, re-verified by hash.",
    "Attested execution: gated commits demand a sealed trust root and a registered UCR instance.",
])
p('Two protocol laws are normative throughout: the Key Identity Law (keys are hashed exactly as given; camelCase and snake_case are distinct inputs) and the Value State Law (absent, null, empty string, empty array, and false are distinct protocol states). Both live in docs/contracts/ORGANISM_RECEIPT_CONTRACT.md.')

h1("Chapter 2 — Organism Receipt Contract (organism_receipt.v1)")
p("One body, one law, one receipt. The unified organism receipt has seven required sections: organ, intent, decision, effect, evidence, replay, continuity.")
code('receipt_id = "org:" + sha256( canonical_json( receipt_with_receipt_id_empty ) )')
p("Canonical form: sorted keys, no whitespace, UTF-8 raw unicode, deterministic re-derivation detects any post-issuance mutation.")
h2("Key Identity Law")
p("Subject and receipt keys are hashed exactly as given. camelCase and snake_case are distinct hash inputs by design; no adapter layer may normalize casing before hashing. Changing normalization is a versioned protocol change (organism_receipt.v2) with dual-issue migration — never a cleanup.")
h2("Value State Law")
p("Absent, null, empty string, empty array, and false are distinct protocol states. Empty optional fields hash their canonical form ('\"\"' hashes as sha256 of two quote characters; '[]' as an empty JSON array) — never collapsed. TS ports must remember undefined is omitted by serialization while null is kept; Python has no undefined, so builders meaning 'absent' must not insert the key at all.")
p("Validator laws: seven sections required; no anonymous actors on accept outcomes; refusals are first-class evidence recording actor_class; outcomes are accept|reject|escalate; dialects are lirl|amul|nx-replay|infinity-trace.")

doc.add_page_break()
h1("Chapter 3 — LIRL Receipt Parity")
p("src/organism_receipt.py maps LIRL stored evidence into organism_receipt.v1. The Sovereign-X Node adapter (src/lirl/organismReceipt.js) is the reference implementation; identical stored evidence must yield identical ids from both runtimes.")
p("A real parity bug was found and fixed during the port: with subjectHash or evidenceRefs absent, the Node adapter hashes the canonical empty value while Python originally collapsed to the empty string — breaking byte-for-byte identity. The fix passes '' through to the digest so Python hashes '\"\"' and '[]' exactly like Node.")
p("Test coverage: four golden vectors (accept, anonymous-refusal, unicode, minimal-empty) with Node-generated golden ids hardcoded, tamper-evidence checks on every vector, unicode canonicalization parity, and a live cross-check against the Sovereign-X working clone when mounted.")

h1("Chapter 4 — Evidence Receipts (AAES Spine)")
p("Sealed claims backing trust bundles: deterministic 'evidence:' receipt ids via sha3-256 over claim_label | subsystem | evidence_refs joined by comma | subject_hash, where the subject hash itself seals the stable-stringified subject.")
code('subject_hash = "sha3-256:" + sha3_256( stable_stringify(subject) )\nreceipt_id   = "evidence:" + sha3_256( claim|subsystem|refs|subject_hash )')
p("Receipt ids are intentionally time-independent: issued_at is provenance metadata, never hashed. Kind inference maps subsystems and claims into fault|patch|mri|trust|attestation|runtime|generic. Specialized sealers exist for CEN enforcement decisions and MRI provenance. The Key Identity Law applies verbatim to subject keys.")

h1("Chapter 5 — Invariant Registry and IDSL-1")
p("Six canonical constitutional invariants with severity and authority-token metadata:")
table = doc.add_table(rows=1, cols=5)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for i, name in enumerate(["ID", "Name", "Dimension", "Floor", "Severity"]): hdr[i].text = name
for inv_id, name, dim, floor, sev in [
    ("INV-003","Governance Drift","governance",70,"high"),
    ("INV-007","Resource Floor","continuity",50,"high"),
    ("INV-014","Temporal Regularity","coordination",55,"medium"),
    ("INV-021","Identity Boundary","memory",60,"critical (VT)"),
    ("INV-031","Coordination Floor","coordination",60,"high"),
    ("INV-041","Confidence Floor","confidence",70,"medium"),
]:
    row = table.add_row().cells
    for i, v in enumerate([inv_id, name, dim, str(floor), sev]): row[i].text = v
p("IDSL-1 compiles boolean invariant expressions without eval: AND/OR/NOT clauses over the five constitutional dimensions (continuity, governance, memory, coordination, confidence), plus legacy 'require <dim> >= <floor>' syntax. Failure profile is deliberately two-stage: unknown dimension words fail the compile-time allowlist; unsupported operators fail when the clause evaluates. Precedence law: transition payload values override mri_snapshot values — conflicts resolve deterministically toward payload, never merged.")

doc.add_page_break()
h1("Chapter 6 — The Constitutional Enforcement Node (EP-1)")
p("Ported from @aaes-os/constitutional-enforcement-node. The governed lifecycle: intercept -> evaluate -> allow/deny, with replay detection, capability gates, invariant evaluation, and hash-chained enforcement receipts.")
code('''decision = { verdict: ALLOW|DENY, action: ALLOW|DENY|FREEZE|MANDATORY_REVIEW,
             reasonCode, reasonDetail }
receipt  = { receiptId: "cen:<hash>", ..., previousReceiptHash, receiptHash }''')
p("Refusal ordering inside evaluate(): malformed transition -> replay detected -> capability denied -> authority-token validation -> invariants. Every finished evaluation appends a chained receipt; verify_enforcement_receipt re-derives the hash from the receipt base. Cross-runtime goldens pin both the tokenless and tokenized receipt branches byte-for-byte against the TS algorithm.")
p("Authority tokens: domain-separated sha3-256 signatures ('AAES-CEN-AUTHORITY-TOKEN-v1'), single-use consumption, validation order: replayed -> invalid signature -> expired -> transition mismatch -> scope denied.")

h1("Chapter 7 — INV-021 and the VT Gate for Law Mutations")
p("The Identity Boundary invariant declares that critical-severity law requires a Verified Token. Enforced rule: any law_mutation transition requires a VT-type authority token. Not a warning, not log-and-continue, not an LLM judgment — denial before commit.")
p("Token validation chain for law mutations: presence -> type must be VT (checked by the bridge before node execution) -> signature -> expiry -> transition binding -> scope contains law:mutate. Refusals still traverse the node so refusal receipts enter the hash chain, carrying the precise reason detail ('law mutations require a VT authority token (INV-021)').")

h1("Chapter 8 — The CEN Governance Bridge")
p("src/cen_governance_bridge.py is the last common admission path before state becomes authoritative. The enforced pipeline:")
code('''proposed mutation -> classify -> CEN.execute() -> INV-021 -> VT verification
  -> enforcement receipt + evidence receipt -> frozen approval -> commit OR deny''')
p("Classification precedence: explicit transition_type argument, then bundle declaration (constitutional_class), else runtime_action. Any exception inside enforcement fails closed — denial with reason cen_failed_closed, commit function never invoked. There is no environment switch that turns the gate off.")

doc.add_page_break()
h1("Chapter 9 — Law-State Sink Dominance")
p('The real invariant is not "the runtime consults CEN" — it is that ALL authoritative law-mutation write paths are dominated by CEN. Dominance lives at the sink: every one of the 14 law registries refuses any save lacking a valid CEN approval envelope bound to the exact record digest.')
table = doc.add_table(rows=1, cols=2); table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells; hdr[0].text = "Sink"; hdr[1].text = "Record"
sinks = ["operator_membrane_policy","autobiographical_episode","social_bond",
         "constitutional_ecosystem_charter","constitutional_evolution_amendment",
         "culture_habit","culture_of_beings_norm","federated_epoch_charter",
         "governed_civilization","identity_self_model_claim","diplomatic_accord",
         "multi_being_pact","narrative_beat","norm_federation_treaty"]
for sname in sinks:
    row = table.add_row().cells; row[0].text = sname; row[1].text = "cen_approval required"
p("Consequence: maintenance scripts, recovery routines, direct state setters, migration jobs, and debug endpoints calling a registry directly hit PermissionError at the sink itself. Even test fixtures route through cen_approved() — fixtures obey the same law as production writers.")
p("Binding refinement: randomly-assigned record ids (pact_id, charter_id, ...) are bookkeeping, not law content — reduction strips each sink's primary id field so challenge-response token minting binds deterministically across attempts. Operational fields policy_id and jarvis_receipt_id are excluded identically.")

h1("Chapter 10 — TOCTOU-Safe Frozen Approvals")
p("Approval binds content, not references. gate_commit deep-copies the payload into a frozen approval stamped with its hash; commit_approved re-verifies the hash immediately before executing the commit function. Mutating the caller's object after approval cannot change what gets committed — verified by tests where post-approval mutation leaves the committed value untouched, and swapped approvals are refused with toctou_hash_mismatch.")
p("WorkflowChainExecutor integration: every non-dry commit passes the gate first; plugs receive the approved frozen args, not caller-held dictionaries. Dry runs never make state authoritative and stay ungated. Denials return blocked runs carrying the CEN decision and receipt ids.")

doc.add_page_break()
h1("Chapter 11 — Mandala Link IMXP Adapter")
p("src/imxp_mandala_adapter.py wraps mandala-link/1 packets into governance membrane events: MGM-0 drift observation, MGM-1 grant-to-policy candidates, and MGM-3 admission-only packet consult. Accepts both wire-spec snake_case and TS-reference camelCase envelopes; canonical hashing matches organism_receipt so receipts stay compatible.")
p("Packet types map onto permeability channels (mesh media to mesh_handoff, text/command/response/control to exchange_envelope, sensor telemetry to memory_cues); channels map onto operator_membrane_policy kinds. Capability grants convert into policy candidates through the dual-gate adoption flow — operator promotion plus Jarvis authorization — which now additionally requires a CEN VT approval at the sink.")

h1("Chapter 12 — Trust Root")
p("src/trust_root.py ports @aaes-os/trust-root: canonical 'sha3-256:' measurement strings validated by strict pattern, and a deterministic hTrustRoot computed domain-separated ('AAES-TRUST-ROOT-v1') over raw measurement bytes in fixed order: kernel image, law spine, corridors, boot manifest.")
code("hTrustRoot = sha3_256( DOMAIN || raw(hKernelImage) || raw(hLawSpine) || raw(hCorridors) || raw(hBootManifest) )")
p("One-shot sealing: the root seals exactly once per process lifetime; re-sealing raises. run_early_boot composes build-and-seal and returns OK. A UCR context projects the law-spine, corridors, and trust-root measurements for downstream attestation.")

h1("Chapter 13 — UCR Attestation")
p("src/ucr_attestation.py ports @aaes-os/ucr-attestation: tokens bind a UCR instance and build fingerprint to the sealed trust root, corridors hash, and law-spine hash, signed with domain 'AAES-UCR-ATTEST-v1'. Signature bytes reproduce the TS Buffer join exactly (domain NUL followed by pipe-separated fields).")
p("Registration refusal ordering is deterministic: unsealed boot -> expired token -> invalid law key (32 hex chars, not all-zero) -> invalid signature -> trust-root mismatch -> corridors mismatch -> law-spine mismatch. Success registers a UCR handle — the identity downstream consumers check.")
p("Pairing into the boundary: gate_commit(require_ucr_attested=True) refuses until the organism has sealed its root and registered an attested instance; refusals chain enforcement and evidence receipts like every other denial. Unsealed -> refused; sealed-but-unregistered -> refused; attested -> commits with frozen-object binding.")

doc.add_page_break()
h1("Chapter 14 — Adversarial Conformance Testing")
p("Cross-runtime goldens are the backbone: receipt ids, trust roots, and attestation signatures generated from the exact TS algorithms via Node and asserted in Python without requiring Node at runtime; live cross-checks run when the Sovereign-X clone is mounted.")
p("The adversarial suites attack the ugly cases:")
bullets([
    "Malformed receipts: every missing section, wrong types, non-object input.",
    "Single-byte mutation detection across all seven receipt sections plus booleans and unicode.",
    "Key reordering produces identical ids; key casing deliberately does not.",
    "Unknown invariants, unknown clauses, eval injection rejected.",
    "Absent / invalid / expired / wrong-scope / wrong-record VT denials for law mutations.",
    "Forged approvals and approvals transplanted to tampered records refused at sinks.",
    "CEN failure fails closed; post-approval mutation refused; exact frozen object committed.",
    "Alternate-path dominance: direct registry writes, forged envelopes, and seed fixtures all dominated.",
])

h1("Chapter 15 — Deployment State, Checkpoints, and Next Steps")
p("Primary working tree (New Volume): f69f79e -> caab105 (enforcement stack + contract laws) -> 79bdbb6 (membrane dominance) -> e06c823 (all 14 sinks) -> bebdf1c (trust-root + attestation). Secondary checkpoint lineage on the DEEA drive carries the synced stack (b9852d7 -> 84bf1f2 -> 6df993a). Upstream GitHub was searched bloblessly: ORGANISM_RECEIPT_CONTRACT.md never existed remotely — the rebuilt contract is canonical.")
table = doc.add_table(rows=1, cols=3); table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for i, name in enumerate(["Commit", "Content", "Location"]): hdr[i].text = name
for cid, content, loc in [
    ("b9852d7", "AAES port with Node-matched goldens", "DEEA clone"),
    ("caab105", "Full stack + ORGANISM_RECEIPT_CONTRACT.md (Key Identity + Value State Laws)", "primary"),
    ("79bdbb6", "Membrane sink dominance + TOCTOU-safe executor wiring", "primary"),
    ("e06c823", "Dominance rolled to all 14 law-state sinks", "primary"),
    ("bebdf1c", "Trust root + UCR attestation feeding CEN", "primary"),
]:
    row = table.add_row().cells
    for i, v in enumerate([cid, content, loc]): row[i].text = v
p("Open threads: push primary to origin (ahead by five commits); sync bebdf1c to the DEEA checkpoint; the Tri-Strata memory spine (runledger, governed-memory) as the next package pair; and the pre-existing test_constitutional_evolution_adopt fixture failure (missing federated epoch registry file) documented in e06c823.")
p("Closing note: constitutional enforcement is now part of the organism's causal execution path rather than something that can merely describe whether an action was constitutional afterward.", bold=True)

doc.save("docs/AAES_CONSTITUTIONAL_ENFORCEMENT_STACK.docx")
print("saved docs/AAES_CONSTITUTIONAL_ENFORCEMENT_STACK.docx")
